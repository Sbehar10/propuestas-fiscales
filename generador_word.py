# generador_word.py — Generador de documentos Word premium (Forbes/Expansión style)
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import io
from datetime import datetime


# Brand colors
NAVY = RGBColor(0x0D, 0x1B, 0x2A)
GOLD = RGBColor(0xC9, 0xA8, 0x4C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
OFF_WHITE = RGBColor(0xF8, 0xF7, 0xF4)
LIGHT_GRAY = RGBColor(0xF3, 0xF4, 0xF6)
BODY_TEXT = RGBColor(0x37, 0x41, 0x51)
SECONDARY = RGBColor(0x6B, 0x72, 0x80)
SUCCESS = RGBColor(0x16, 0xA3, 0x4A)
DARK_GREEN = RGBColor(0x15, 0x80, 0x3D)

# Legacy aliases for backward compat
AZUL_OSCURO = NAVY
DORADO = GOLD
BLANCO = WHITE
GRIS_CLARO = LIGHT_GRAY
VERDE = SUCCESS
NEGRO = BODY_TEXT


def fmt_moneda(valor):
    """Formatea número como moneda mexicana"""
    if valor is None:
        return "$0.00"
    return f"${valor:,.2f}"


def fmt_pct(valor):
    """Formatea porcentaje"""
    return f"{valor:.2f}%"


def set_cell_shading(cell, color_hex):
    """Aplica sombreado a celda"""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elem)


def _set_cell_border(cell, **kwargs):
    """Set cell borders. Usage: _set_cell_border(cell, top={"sz":4,"color":"C9A84C"})"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="single" '
            f'w:sz="{attrs.get("sz", 4)}" w:space="0" '
            f'w:color="{attrs.get("color", "0D1B2A")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def _add_gold_separator(doc):
    """Adds a thin gold line separator"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 60)
    run.font.color.rgb = GOLD
    run.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)


def _add_section_heading(doc, text, number=None):
    """Adds a premium section heading with gold underline"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    label = f"{number}. {text}" if number else text
    run = p.add_run(label)
    run.font.name = 'Cambria'
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY
    run.bold = True

    # Gold underline
    p2 = doc.add_paragraph()
    run2 = p2.add_run("━" * 50)
    run2.font.color.rgb = GOLD
    run2.font.size = Pt(6)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(10)


def _add_body_text(doc, text):
    """Adds body text in Calibri with proper color"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.color.rgb = BODY_TEXT
    p.paragraph_format.space_after = Pt(8)
    return p


def _add_kpi_table(doc, kpis):
    """
    Adds KPI boxes in a row. kpis = list of (label, value, is_highlight) tuples.
    """
    table = doc.add_table(rows=2, cols=len(kpis))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (label, value, highlight) in enumerate(kpis):
        # Header cell (label)
        cell_top = table.rows[0].cells[i]
        cell_top.text = ""
        set_cell_shading(cell_top, "0D1B2A")
        p = cell_top.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label.upper())
        run.font.name = 'Calibri'
        run.font.size = Pt(7.5)
        run.font.color.rgb = WHITE
        run.bold = True

        # Value cell
        cell_bot = table.rows[1].cells[i]
        cell_bot.text = ""
        set_cell_shading(cell_bot, "FFFFFF")
        p2 = cell_bot.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(str(value))
        run2.font.name = 'Calibri'
        run2.font.size = Pt(14)
        run2.bold = True
        run2.font.color.rgb = SUCCESS if highlight else NAVY

    # Remove default table borders and add subtle ones
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(cell,
                top={"sz": 4, "color": "0D1B2A"},
                bottom={"sz": 4, "color": "0D1B2A"},
                left={"sz": 4, "color": "E5E3DB"},
                right={"sz": 4, "color": "E5E3DB"})

    doc.add_paragraph("")  # spacing


def agregar_tabla_estilizada(doc, headers, rows, col_widths=None):
    """Crea tabla profesional con estilo premium"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row — navy background, white text
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(8.5)
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "0D1B2A")

    # Data rows — alternating white / light gray
    for r_idx, row_data in enumerate(rows):
        for c_idx, value in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(value))
            run.font.size = Pt(8.5)
            run.font.name = 'Calibri'
            run.font.color.rgb = BODY_TEXT
            if c_idx > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F3F4F6")

    # Bold TOTAL row
    if rows and "TOTAL" in str(rows[-1][0]).upper():
        for c_idx in range(len(headers)):
            cell = table.rows[len(rows)].cells[c_idx]
            set_cell_shading(cell, "0D1B2A")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = WHITE

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def generar_propuesta_word(datos_cliente, tipo_servicio, resultados, grupos=None):
    """
    Genera documento Word premium con la propuesta completa.
    """
    doc = Document()

    # Configurar márgenes
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # === PORTADA ===
    _generar_portada(doc, datos_cliente, tipo_servicio)

    # === RESUMEN EJECUTIVO ===
    doc.add_page_break()
    _generar_resumen_ejecutivo(doc, datos_cliente, tipo_servicio, resultados)

    # === CONTENIDO SEGÚN TIPO DE SERVICIO ===
    if tipo_servicio == "nomina" and grupos:
        _generar_seccion_diagnostico(doc, grupos)
        _generar_seccion_propuesta_irt(doc, grupos)
        _generar_seccion_ahorro(doc, resultados, tipo_servicio)
        _generar_seccion_facturacion(doc, datos_cliente, resultados, tipo_servicio)

    elif tipo_servicio == "excedentes":
        _generar_seccion_excedentes(doc, resultados)

    elif tipo_servicio == "sc":
        _generar_seccion_sc(doc, resultados)

    # === FUNDAMENTO LEGAL ===
    doc.add_page_break()
    _generar_fundamento_legal(doc, tipo_servicio)

    # === PLAN DE IMPLEMENTACIÓN ===
    _generar_plan_implementacion(doc, tipo_servicio)

    # === FOOTER ===
    _agregar_footer(doc, datos_cliente["nombre_empresa"])

    # Guardar a buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _generar_portada(doc, datos_cliente, tipo_servicio):
    """Premium cover page — navy top block, gold accents"""
    # Navy header block via table (full width, colored background)
    cover_table = doc.add_table(rows=1, cols=1)
    cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = cover_table.rows[0].cells[0]
    set_cell_shading(cell, "0D1B2A")

    # Build cover content inside the cell
    cell.text = ""

    # Spacer
    p_space = cell.paragraphs[0]
    p_space.text = ""
    p_space.paragraph_format.space_before = Pt(40)

    # Brand name
    p_brand = cell.add_paragraph()
    p_brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_brand.paragraph_format.space_before = Pt(10)
    run = p_brand.add_run("")
    run.font.name = 'Cambria'
    run.font.size = Pt(12)
    run.font.color.rgb = GOLD
    run.bold = True

    # Gold separator
    p_line = cell.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_line.add_run("━" * 40)
    run.font.color.rgb = GOLD
    run.font.size = Pt(8)

    # Main title
    titulos = {
        "nomina": "PROPUESTA DE\nOPTIMIZACIÓN FISCAL",
        "excedentes": "PROPUESTA DE\nADMINISTRACIÓN DE EXCEDENTES",
        "sc": "PROPUESTA DE\nSOCIEDAD CIVIL",
    }
    p_title = cell.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(16)
    run = p_title.add_run(titulos.get(tipo_servicio, "PROPUESTA DE SERVICIOS"))
    run.font.name = 'Cambria'
    run.font.size = Pt(28)
    run.font.color.rgb = WHITE
    run.bold = True

    # Year
    p_year = cell.add_paragraph()
    p_year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_year.paragraph_format.space_before = Pt(8)
    run = p_year.add_run("2026")
    run.font.name = 'Cambria'
    run.font.size = Pt(36)
    run.font.color.rgb = GOLD
    run.bold = True

    # Gold separator
    p_line2 = cell.add_paragraph()
    p_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_line2.add_run("━" * 40)
    run.font.color.rgb = GOLD
    run.font.size = Pt(8)

    # Client name
    p_client = cell.add_paragraph()
    p_client.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_client.paragraph_format.space_before = Pt(14)
    run = p_client.add_run(f"Preparada para")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    p_name = cell.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_name.add_run(datos_cliente['nombre_empresa'])
    run.font.name = 'Cambria'
    run.font.size = Pt(18)
    run.font.color.rgb = WHITE
    run.bold = True

    # Contact
    p_contact = cell.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_before = Pt(4)
    run = p_contact.add_run(f"Atención: {datos_cliente['contacto']}")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    # Bottom spacer + brand tagline
    p_tag = cell.add_paragraph()
    p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tag.paragraph_format.space_before = Pt(30)
    p_tag.paragraph_format.space_after = Pt(20)
    run = p_tag.add_run("Consultoría Fiscal Estratégica")
    run.font.name = 'Calibri'
    run.font.size = Pt(8)
    run.font.color.rgb = GOLD

    # Date + confidential below the navy block
    doc.add_paragraph("")
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_date.add_run(datetime.now().strftime("%d de %B de %Y"))
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = SECONDARY

    p_conf = doc.add_paragraph()
    p_conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_conf.add_run("CONFIDENCIAL")
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    run.font.color.rgb = GOLD
    run.bold = True


def _generar_resumen_ejecutivo(doc, datos_cliente, tipo_servicio, resultados):
    """Sección de resumen ejecutivo con KPI boxes"""
    _add_section_heading(doc, "Resumen Ejecutivo", 1)

    if tipo_servicio == "nomina":
        total_empleados = resultados.get("total_empleados", 0)
        ahorro_mensual = resultados.get("ahorro_total_mensual", 0)
        costo_actual = resultados.get("costo_actual_total", 0)
        es_neto = resultados.get("es_neto", False)

        _add_body_text(doc,
            f"La presente propuesta contempla la administración de nómina para "
            f"{total_empleados} colaboradores mediante un esquema de servicios especializados "
            f"con REPSE vigente, optimizando la carga fiscal y patronal de su empresa."
        )

        if es_neto:
            p = doc.add_paragraph()
            run = p.add_run(
                "Sus empleados siguen recibiendo el mismo neto. "
                "Lo que cambia es el costo patronal."
            )
            run.bold = True
            run.font.color.rgb = SUCCESS
            run.font.name = 'Calibri'

        # KPI boxes
        _add_kpi_table(doc, [
            ("Empleados", str(total_empleados), False),
            ("Ahorro Mensual", fmt_moneda(ahorro_mensual), True),
            ("Ahorro Anual", fmt_moneda(ahorro_mensual * 12), True),
            ("Ahorro 3 Años", fmt_moneda(ahorro_mensual * 36), True),
        ])

    elif tipo_servicio == "excedentes":
        _add_body_text(doc,
            f"La presente propuesta contempla la administración de excedentes salariales "
            f"(bonos, comisiones, viáticos) a través de un esquema exento, generando un "
            f"ahorro significativo para su empresa."
        )
        ahorro = resultados.get("ahorro_mensual", 0)
        _add_kpi_table(doc, [
            ("Excedente", fmt_moneda(resultados.get("monto_excedente", 0)), False),
            ("Ahorro Mensual", fmt_moneda(ahorro), True),
            ("Ahorro Anual", fmt_moneda(ahorro * 12), True),
        ])

    elif tipo_servicio == "sc":
        _add_body_text(doc,
            f"La presente propuesta contempla la optimización fiscal para perfiles directivos "
            f"y gerenciales a través de una Sociedad Civil, donde el ingreso se divide entre "
            f"anticipo por remanente (gravado) y renta vitalicia (exenta al 100%)."
        )
        if isinstance(resultados, list):
            for r in resultados:
                _tabla_resumen_sc(doc, r)
        else:
            _tabla_resumen_sc(doc, resultados)


def _tabla_resumen_sc(doc, r):
    """Mini tabla resumen para un directivo SC"""
    headers = ["Concepto", "Esquema Actual (Nómina)", "Esquema SC"]
    rows = [
        ["Ingreso total", fmt_moneda(r["ingreso_total"]), fmt_moneda(r["ingreso_total"])],
        ["ISR retenido", fmt_moneda(r["isr_nomina"]["isr_neto"]), fmt_moneda(r["isr_anticipo"]["isr_neto"])],
        ["Neto al directivo", fmt_moneda(r["neto_nomina"]), fmt_moneda(r["neto_total"])],
        ["Costo total empresa", fmt_moneda(r["costo_nomina_100"]), fmt_moneda(r["total_factura"])],
        ["Ahorro mensual", "—", fmt_moneda(r["ahorro_cliente_mensual"])],
    ]
    agregar_tabla_estilizada(doc, headers, rows)
    doc.add_paragraph("")


def _generar_seccion_diagnostico(doc, grupos):
    """Sección 2: Diagnóstico fiscal actual"""
    doc.add_page_break()
    _add_section_heading(doc, "Diagnóstico Fiscal Actual", 2)

    _add_body_text(doc,
        "A continuación se presenta el costo actual que representa la nómina "
        "de los colaboradores bajo un esquema 100% formal ante el IMSS."
    )

    headers = ["Puesto", "Cant.", "Sueldo Bruto (c/u)", "ISR", "IMSS Pat.", "IMSS Obr.",
               "Infonavit", "ISN", "Prestaciones", "Costo Total/Mes"]
    rows = []
    totales = {"isr": 0, "imss_pat": 0, "imss_obr": 0, "info": 0, "isn": 0, "prest": 0, "costo": 0}

    for g in grupos:
        a = g["actual"]
        n = g["num_empleados"]
        row = [
            g["puesto"],
            str(n),
            fmt_moneda(a["sueldo_bruto"]),
            fmt_moneda(a["isr"]["isr_neto"] * n),
            fmt_moneda(a["imss_patronal"]["total"] * n),
            fmt_moneda(a["imss_obrero"]["total"] * n),
            fmt_moneda(a["infonavit"] * n),
            fmt_moneda(a["isn"] * n),
            fmt_moneda(a["prestaciones"]["total_mensual"] * n),
            fmt_moneda(a["costo_total"]),
        ]
        rows.append(row)
        totales["isr"] += a["isr"]["isr_neto"] * n
        totales["imss_pat"] += a["imss_patronal"]["total"] * n
        totales["imss_obr"] += a["imss_obrero"]["total"] * n
        totales["info"] += a["infonavit"] * n
        totales["isn"] += a["isn"] * n
        totales["prest"] += a["prestaciones"]["total_mensual"] * n
        totales["costo"] += a["costo_total"]

    rows.append([
        "TOTAL", "", "",
        fmt_moneda(totales["isr"]),
        fmt_moneda(totales["imss_pat"]),
        fmt_moneda(totales["imss_obr"]),
        fmt_moneda(totales["info"]),
        fmt_moneda(totales["isn"]),
        fmt_moneda(totales["prest"]),
        fmt_moneda(totales["costo"]),
    ])

    agregar_tabla_estilizada(doc, headers, rows)


def _generar_seccion_propuesta_irt(doc, grupos):
    """Sección 3: Propuesta IRT con comparativo Actual vs IRT"""
    doc.add_page_break()
    _add_section_heading(doc, "Propuesta de Optimización — IRT", 3)

    _add_body_text(doc,
        "Se propone optimizar la carga fiscal y patronal mediante un esquema de "
        "Indemnización por Riesgo de Trabajo (IRT), donde se cotiza al colaborador "
        "con el salario mínimo profesional correspondiente a su puesto según la tabla "
        "del IMSS, y el excedente se paga como concepto exento conforme al Art. 93 "
        "fracción III de la LISR."
    )

    for g in grupos:
        neto_orig = g.get("sueldo_neto_original")
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run(f"{g['puesto']} ({g['num_empleados']} empleados)")
        run.font.name = 'Cambria'
        run.font.size = Pt(12)
        run.font.color.rgb = NAVY
        run.bold = True

        headers = ["Concepto", "Actual (100%)", "IRT Propuesto"]
        a = g["actual"]
        irt = g["irt"]

        rows = []
        if neto_orig:
            rows.append(["Sueldo neto del empleado", fmt_moneda(neto_orig), fmt_moneda(neto_orig) + " (sin cambio)"])
            rows.append(["Bruto equivalente", fmt_moneda(a["sueldo_bruto"]), fmt_moneda(a["sueldo_bruto"])])
        else:
            rows.append(["Sueldo bruto", fmt_moneda(a["sueldo_bruto"]), fmt_moneda(a["sueldo_bruto"])])
        rows += [
            ["Base nómina (IMSS)", fmt_moneda(a["sueldo_bruto"]), fmt_moneda(irt["base_nomina"])],
            ["Excedente IRT (exento)", "—", fmt_moneda(irt["excedente_irt"])],
            ["ISR mensual", fmt_moneda(a["isr"]["isr_neto"]), fmt_moneda(irt["isr"]["isr_neto"])],
            ["IMSS patronal", fmt_moneda(a["imss_patronal"]["total"]), fmt_moneda(irt["imss_patronal"]["total"])],
            ["IMSS obrero", fmt_moneda(a["imss_obrero"]["total"]), fmt_moneda(irt["imss_obrero"]["total"])],
            ["Infonavit", fmt_moneda(a["infonavit"]), fmt_moneda(irt["infonavit"])],
            ["ISN", fmt_moneda(a["isn"]), fmt_moneda(irt["isn"])],
            ["Neto trabajador", fmt_moneda(a["neto_trabajador"]), fmt_moneda(irt["neto_trabajador"])],
        ]
        agregar_tabla_estilizada(doc, headers, rows)
        doc.add_paragraph("")


def _generar_seccion_ahorro(doc, resultados, tipo_servicio):
    """Sección 4: Análisis de ahorro"""
    doc.add_page_break()
    _add_section_heading(doc, "Análisis de Ahorro", 4)

    costo_actual = resultados.get("costo_actual_total", 0)
    ahorro = resultados.get("ahorro_total_mensual", 0)
    pct = (ahorro / costo_actual * 100) if costo_actual > 0 else 0

    # KPI highlight
    _add_kpi_table(doc, [
        ("Costo Actual", fmt_moneda(costo_actual), False),
        ("Ahorro Mensual", fmt_moneda(ahorro), True),
        ("% Ahorro", fmt_pct(pct), True),
    ])

    headers = ["Período", "Ahorro IRT", "% Ahorro"]
    rows = [
        ["Mensual", fmt_moneda(ahorro), fmt_pct(pct)],
        ["Anual", fmt_moneda(ahorro * 12), fmt_pct(pct)],
        ["2 años", fmt_moneda(ahorro * 24), fmt_pct(pct)],
        ["3 años", fmt_moneda(ahorro * 36), fmt_pct(pct)],
    ]
    agregar_tabla_estilizada(doc, headers, rows)

    # Highlighted callout
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run(f"Ahorro proyectado a 3 años: {fmt_moneda(ahorro * 36)}")
    run.font.size = Pt(14)
    run.font.color.rgb = SUCCESS
    run.font.name = 'Cambria'
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _generar_seccion_facturacion(doc, datos_cliente, resultados, tipo_servicio):
    """Sección 5: Detalle de facturación"""
    _add_section_heading(doc, "Detalle de Facturación", 5)

    _add_body_text(doc,
        "La facturación mensual se compone de la siguiente manera:"
    )

    comision_pct = datos_cliente["comision_pct"]
    total_admin = resultados.get("total_administrado", 0)
    comision = resultados.get("total_comision", round(total_admin * (comision_pct / 100), 2))
    subtotal = resultados.get("subtotal_factura", total_admin + comision)
    iva = resultados.get("iva_total", round(subtotal * 0.16, 2))
    total = resultados.get("total_factura", subtotal + iva)

    headers = ["Concepto", "Monto"]
    rows = [
        ["Nómina total administrada", fmt_moneda(total_admin)],
        [f"Comisión ({comision_pct}%)", fmt_moneda(comision)],
        ["Subtotal", fmt_moneda(subtotal)],
        ["IVA (16%)", fmt_moneda(iva)],
        ["TOTAL FACTURA MENSUAL", fmt_moneda(total)],
    ]
    agregar_tabla_estilizada(doc, headers, rows)

    doc.add_paragraph("")
    _add_body_text(doc,
        "El gasto total es 100% deducible para efectos del ISR de su empresa, "
        "ya que se trata de un servicio especializado facturado con IVA y CFDI vigente."
    )


def _generar_seccion_excedentes(doc, resultados):
    """Sección específica para excedentes"""
    _add_section_heading(doc, "Análisis de Excedentes", 2)

    headers = ["Concepto", "Monto"]
    rows = [
        ["Excedente mensual", fmt_moneda(resultados["monto_excedente"])],
        ["Comisión", fmt_moneda(resultados["comision"])],
        ["IVA", fmt_moneda(resultados["iva"])],
        ["Total factura mensual", fmt_moneda(resultados["total_factura"])],
        ["", ""],
        ["Costo si pagara por nómina", fmt_moneda(resultados["costo_hipotetico_nomina"])],
        ["Ahorro mensual", fmt_moneda(resultados["ahorro_mensual"])],
        ["Ahorro anual", fmt_moneda(resultados["ahorro_anual"])],
    ]
    agregar_tabla_estilizada(doc, headers, rows)


def _generar_seccion_sc(doc, resultados):
    """Sección específica para Sociedad Civil"""
    _add_section_heading(doc, "Análisis Sociedad Civil", 2)

    _add_body_text(doc,
        "El ingreso del directivo/socio se estructura de la siguiente manera a través "
        "de la Sociedad Civil, dividiendo entre anticipo por remanente (gravado, con "
        "retención ISR Art. 96) y renta vitalicia (exenta al 100%, Art. 93 fr. IV LISR)."
    )

    lista = resultados if isinstance(resultados, list) else [resultados]
    for r in lista:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run(f"Ingreso: {fmt_moneda(r['ingreso_total'])} — Anticipo {r['pct_anticipo']}%")
        run.font.name = 'Cambria'
        run.font.size = Pt(12)
        run.font.color.rgb = NAVY
        run.bold = True

        headers = ["Concepto", "Monto"]
        rows = [
            ["Ingreso total", fmt_moneda(r["ingreso_total"])],
            [f"Anticipo por remanente ({r['pct_anticipo']}%)", fmt_moneda(r["anticipo"])],
            ["ISR sobre anticipo", fmt_moneda(r["isr_anticipo"]["isr_neto"])],
            ["Renta vitalicia (exenta)", fmt_moneda(r["renta"])],
            ["Neto al directivo", fmt_moneda(r["neto_total"])],
            ["", ""],
            ["Comisión", fmt_moneda(r["comision"])],
            ["IVA", fmt_moneda(r["iva"])],
            ["Total factura", fmt_moneda(r["total_factura"])],
            ["", ""],
            ["VS Nómina 100% — Costo empresa", fmt_moneda(r["costo_nomina_100"])],
            ["VS Nómina 100% — Neto directivo", fmt_moneda(r["neto_nomina"])],
            ["Ahorro mensual", fmt_moneda(r["ahorro_cliente_mensual"])],
            ["Ahorro anual", fmt_moneda(r["ahorro_cliente_anual"])],
        ]
        agregar_tabla_estilizada(doc, headers, rows)
        doc.add_paragraph("")


def _generar_fundamento_legal(doc, tipo_servicio):
    """Sección de fundamento legal"""
    _add_section_heading(doc, "Fundamento Legal")

    if tipo_servicio in ("nomina", "excedentes"):
        _add_body_text(doc,
            "El presente esquema se sustenta en los siguientes ordenamientos legales:"
        )
        fundamentos = [
            "Art. 93 fracción III LISR — Exención de indemnizaciones por riesgos o enfermedades.",
            "Art. 93 fracción XIII LISR — Exención de viáticos efectivamente erogados.",
            "NOM-035-STPS-2018 — Factores de riesgo psicosocial en el trabajo.",
            "Art. 15-A Ley Federal del Trabajo — Servicios especializados.",
            "Art. 13 Ley del Seguro Social — Obligaciones del patrón.",
            "Registro REPSE vigente ante la STPS.",
        ]
        for f in fundamentos:
            p = doc.add_paragraph()
            run = p.add_run(f"    •  {f}")
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.color.rgb = BODY_TEXT

    elif tipo_servicio == "sc":
        _add_body_text(doc,
            "El esquema de Sociedad Civil se fundamenta en:"
        )
        fundamentos = [
            "Art. 94 fracción II LISR — Ingresos asimilados a salarios por anticipos de remanente.",
            "Art. 96 LISR — Tabla de retención de ISR a personas físicas.",
            "Art. 93 fracción IV LISR — Exención de rentas vitalicias.",
            "Art. 86 LISR — Obligaciones de personas morales con fines no lucrativos.",
            "Código Civil Federal — Constitución y operación de Sociedades Civiles.",
        ]
        for f in fundamentos:
            p = doc.add_paragraph()
            run = p.add_run(f"    •  {f}")
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.color.rgb = BODY_TEXT


def _generar_plan_implementacion(doc, tipo_servicio):
    """Plan de implementación por fases"""
    doc.add_paragraph("")
    _add_section_heading(doc, "Plan de Implementación")

    fases = [
        ("Fase 1 — Diagnóstico (Semana 1)", "Recepción de nómina actual, análisis de puestos, asignación de salarios mínimos profesionales y segmentación por actividad."),
        ("Fase 2 — Estructuración (Semana 2)", "Alta de personal en empresas con REPSE correspondiente, configuración de esquema de pago y validación de cálculos."),
        ("Fase 3 — Migración (Semana 3-4)", "Transición de colaboradores, primer dispersión de nómina bajo nuevo esquema, validación de CFDIs."),
        ("Fase 4 — Operación continua", "Dispersión quincenal/mensual, reportes de nómina, conciliaciones, cumplimiento de obligaciones fiscales y laborales."),
    ]

    for titulo, descripcion in fases:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        run = p.add_run(titulo)
        run.font.name = 'Cambria'
        run.font.size = Pt(11)
        run.font.color.rgb = GOLD
        run.bold = True
        _add_body_text(doc, descripcion)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Requisitos documentales del cliente:")
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.color.rgb = NAVY
    run.bold = True

    reqs = [
        "Nómina actual completa (puestos, sueldos, prestaciones)",
        "Comprobante de inscripción fiscal (CIF)",
        "Organigrama o descripción de puestos",
        "Contrato de servicios especializado firmado",
    ]
    for r in reqs:
        p = doc.add_paragraph()
        run = p.add_run(f"    •  {r}")
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = BODY_TEXT


def _agregar_footer(doc, nombre_empresa):
    """Footer: gold line + confidential text"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False

        # Gold separator line
        p_line = footer.paragraphs[0]
        p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_line.add_run("━" * 50)
        run.font.color.rgb = GOLD
        run.font.size = Pt(6)

        # Confidential text
        p = footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Confidencial — Propuesta Fiscal 2026  |  Propuesta para {nombre_empresa}")
        run.font.size = Pt(7)
        run.font.color.rgb = SECONDARY
        run.font.name = 'Calibri'
